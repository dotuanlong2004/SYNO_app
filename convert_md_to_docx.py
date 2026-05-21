from docx import Document
from docx.shared import Pt
import sys
import os
import re

FILES = [
    ("1_YEU_CAU.md", "1_YEU_CAU.docx"),
    ("2_INPUT_OUTPUT.md", "2_INPUT_OUTPUT.docx"),
    ("3_THIET_KE_HE_THONG.md", "3_THIET_KE_HE_THONG.docx"),
]

def add_paragraph_for_line(doc, line, in_code):
    line = line.rstrip("\n")
    if in_code:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        return

    # Headings: #, ##, ### ...
    m = re.match(r'^(#{1,6})\\s+(.*)$', line)
    if m:
        level = len(m.group(1))
        text = m.group(2)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        # approximate sizes
        sizes = {1:18, 2:16, 3:14, 4:12, 5:11, 6:10}
        run.font.size = Pt(sizes.get(level, 11))
        return

    # Horizontal rule
    if re.match(r'^(-{3,}|\*{3,})\s*$', line):
        p = doc.add_paragraph()
        run = p.add_run('--------------------------------------------------')
        return

    # List items (unordered or ordered)
    m = re.match(r'^\\s*([-\\*\\+]\\s+)(.*)$', line)
    if m:
        p = doc.add_paragraph(m.group(2))
        p.style = 'List Bullet'
        return
    m = re.match(r'^\\s*(\\d+\\.\\s+)(.*)$', line)
    if m:
        p = doc.add_paragraph(m.group(2))
        p.style = 'List Number'
        return

    # Tables: detect pipe-separated header or row -> add as plain paragraph
    if '|' in line and not line.strip().startswith('```'):
        # keep table lines as-is (simple)
        p = doc.add_paragraph(line)
        return

    # Blank line -> add paragraph break
    if line.strip() == '':
        doc.add_paragraph()
        return

    # Normal paragraph
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(11)

def convert(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"SKIP: {md_path} not found")
        return False
    doc = Document()
    in_code = False
    with open(md_path, 'r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\\n')
            # handle fenced code blocks
            if line.strip().startswith("```"):
                in_code = not in_code
                if in_code:
                    # start code block - optionally add a separator
                    doc.add_paragraph()
                else:
                    doc.add_paragraph()
                continue
            add_paragraph_for_line(doc, raw, in_code)
    doc.save(docx_path)
    print(f"CREATED: {docx_path}")
    return True

def main():
    success = True
    for md, docx in FILES:
        ok = convert(md, docx)
        success = success and ok
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()